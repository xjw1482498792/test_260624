# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = "简历_肖家伟_SAP_ABAP_v2_优化版.docx"
FONT = "Microsoft YaHei"
BLUE = RGBColor(0x05, 0x63, 0xC7)
DARK_BLUE = RGBColor(0x0B, 0x4F, 0x8A)
TEXT = RGBColor(0x2D, 0x2D, 0x2D)
GRAY = RGBColor(0x6A, 0x6A, 0x6A)


def font(run, size=8.7, bold=False, color=TEXT):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=0, start=80, bottom=0, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_layout(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def no_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        borders.append(e)


def paragraph(cell_or_doc, before=0, after=0, line=1.0, left=0, hanging=0):
    p = cell_or_doc.add_paragraph() if hasattr(cell_or_doc, "add_paragraph") else None
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if left:
        p.paragraph_format.left_indent = Cm(left)
    if hanging:
        p.paragraph_format.first_line_indent = Cm(-hanging)
    return p


def add_section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3.8)
    p.paragraph_format.space_after = Pt(2.2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    font(r, 11.5, True, BLUE)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "0563C7")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_company(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.keep_with_next = True
    font(p.add_run("松下信息系统（上海）有限公司 大连分公司"), 10.1, True, DARK_BLUE)
    font(p.add_run("  |  SAP ABAP 开发工程师  |  2024.07 - 至今"), 8.8, False, GRAY)


def add_project(doc, name, client, period, desc, work, tech):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1.8)
    p.paragraph_format.space_after = Pt(0.7)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(name), 9.2, True, BLUE)
    font(p.add_run(f"  |  {client}  |  {period}"), 7.9, False, GRAY)

    for label, value in (("项目场景", desc), ("负责内容", work), ("技术重点", tech)):
        q = doc.add_paragraph()
        q.paragraph_format.left_indent = Cm(0.35)
        q.paragraph_format.space_after = Pt(0.35)
        q.paragraph_format.line_spacing = 1.0
        q.paragraph_format.keep_together = True
        font(q.add_run(f"{label}："), 8.15, True)
        font(q.add_run(value), 8.15)


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.43)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(0.7)
    p.paragraph_format.line_spacing = 1.02
    font(p.add_run("• "), 8.2, True, BLUE)
    font(p.add_run(text), 8.15)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    # The one-page resume is intentionally centered vertically by balancing
    # the natural whitespace above and below the content block.
    sec.top_margin = Cm(2.45)
    sec.bottom_margin = Cm(1.05)
    sec.left_margin = Cm(1.55)
    sec.right_margin = Cm(1.55)
    sec.header_distance = Cm(0.5)
    sec.footer_distance = Cm(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(8.7)
    normal.paragraph_format.space_after = Pt(0)

    # Header: name + target role
    t = doc.add_table(rows=1, cols=2)
    no_table_borders(t)
    set_repeat_table_layout(t, [2450, 7780])
    for c in t.rows[0].cells:
        set_cell_margins(c, 0, 0, 0, 0)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    p = t.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    font(p.add_run("肖家伟"), 22, True, DARK_BLUE)
    p = t.cell(0, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    font(p.add_run("SAP ABAP 开发工程师"), 12, True, BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2.5)
    font(p.add_run("15841125092  |  1482498792@qq.com  |  大连（在职）  |  CET-4"), 8.5, False, GRAY)

    add_section(doc, "教育背景")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    font(p.add_run("大连民族大学"), 9.2, True, DARK_BLUE)
    font(p.add_run("  |  软件工程（本科）  |  2020.09 - 2024.06"), 8.4, False, GRAY)

    add_section(doc, "工作及项目经历")
    add_company(doc)
    add_project(doc, "PAVCX S/4HANA 导入企划项目", "厦门松下电子信息有限公司", "2024.10 - 2025.05",
                "R/3旧系统向S/4HANA迁移，并结合新业务需求完成功能调整。",
                "负责14个自开发程序及7个共通程序的需求分析、逻辑新增、单元测试与缺陷修正。",
                "Report报表、RFC接口、Open SQL及S/4HANA适配。")
    add_project(doc, "PMRZ ECC 升级项目", "珠海松下马达有限公司", "2025.06 - 2025.07",
                "SAP ECC升级至S/4HANA，完成既有自开发程序兼容性改造。",
                "负责6个程序迁移，排查过时语法与结构差异，完成代码调整、测试及性能检查。",
                "S/4HANA兼容性改造、ABAP语法调整、Open SQL性能优化。")
    add_project(doc, "PIDQD SCM 革新项目", "青岛松下电子部品（保税区）有限公司", "2025.08 - 2025.11",
                "Galileo旧系统业务与数据迁移至SAP S/4HANA，并承接新增需求。",
                "负责10个自开发程序、5个共通程序改修及4处增强实施，覆盖开发、测试与传输交付。",
                "Report、RFC、Enhancement、SmartForms及接口/打印场景。")
    add_project(doc, "PEACCN 系统构筑项目", "松下娱乐互动（中国）有限公司", "2025.11 - 2026.03",
                "新工厂上线S/4HANA，基于既有工厂程序扩展多工厂共用逻辑。",
                "负责25个共通程序改修；兼任开发担当，统筹任务分配、工时跟踪、传输清单与版本检查。",
                "多工厂逻辑复用、影响范围分析、回归测试及DEV-QAS-PRD传输管理。")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1.8)
    p.paragraph_format.space_after = Pt(0.5)
    font(p.add_run("多工厂日常运维与需求开发"), 9.2, True, BLUE)
    q = doc.add_paragraph()
    q.paragraph_format.left_indent = Cm(0.35)
    q.paragraph_format.space_after = Pt(0.5)
    q.paragraph_format.line_spacing = 1.0
    font(q.add_run("持续处理各工厂运维Ticket与新增需求，完成问题定位、程序改修、测试验证及上线跟踪。"), 8.15)

    add_section(doc, "专业技能")
    skills = [
        "近2年SAP ABAP开发经验，参与4个S/4HANA导入、升级及构筑项目，覆盖需求理解、开发、测试、传输与上线运维。",
        "熟练使用Report、ALV、RFC、BDC、SmartForms及Enhancement；熟悉Open SQL、内表处理与常见性能优化方法，可使用ST05辅助定位SQL问题。",
        "熟练使用断点、消息/语句/异常断点及Watchpoint定位问题；熟悉DEV-QAS-PRD传输流程、请求依赖检查与版本核对。",
        "具备CDS View、OData、RAP及Fiori Elements基础知识；了解S/4HANA 2023及ABAP 7.58相关开发特性。",
        "熟悉SU01用户维护、PFCG角色权限配置、SM30表维护、SM36/SM37后台Job调度与监控，以及SQ01 Query开发。",
        "了解SD、MM、PP、FICO核心业务流程及常用表结构，能够配合功能顾问完成需求澄清与问题定位。",
    ]
    for s in skills:
        add_bullet(doc, s)

    # Clean metadata
    doc.core_properties.title = "肖家伟 - SAP ABAP 开发工程师简历"
    doc.core_properties.subject = "SAP ABAP 求职简历"
    doc.core_properties.author = "肖家伟"
    doc.core_properties.keywords = "SAP, ABAP, S4HANA"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
