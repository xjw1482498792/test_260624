from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = "肖家伟_AI应用开发简历_仿引用样式.docx"
FONT = "Microsoft YaHei"
BLUE = RGBColor(44, 90, 160)
TEXT = RGBColor(34, 34, 34)


def set_run(run, size=9.2, bold=False, color=TEXT):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_center(doc, text, size=10, bold=False, color=TEXT, after=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    return p


def add_header_line(doc, text, size=10, bold=False, color=TEXT, after=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    return p


def add_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=12, bold=True, color=BLUE)

    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "2C5AA0")
    border.append(bottom)
    p_pr.append(border)


def add_line(doc, text="", size=9.2, bold=False, after=1.2, left=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    if left:
        p.paragraph_format.left_indent = Cm(left)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.first_line_indent = Cm(-0.18)
    p.paragraph_format.space_after = Pt(1.1)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run("- " + text)
    set_run(r, size=9.2)


def add_role_line(doc, left, right):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(left)
    set_run(r1, size=10, bold=True)
    r2 = p.add_run("    " + right)
    set_run(r2, size=9.2, bold=True)


def add_project(doc, name, meta, desc, duties, tech, result):
    add_role_line(doc, name, meta)
    add_bullet(doc, "项目描述：" + desc)
    for duty in duties:
        add_bullet(doc, duty)
    add_bullet(doc, "技术栈：" + tech)
    add_bullet(doc, "项目成果：" + result)


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9.2)
    return doc


def build():
    doc = setup_doc()

    add_header_line(doc, "肖家伟", size=22, bold=True, color=RGBColor(26, 26, 26), after=1)
    add_header_line(doc, "AI 应用开发工程师", size=11.5, bold=True, after=1)
    add_header_line(doc, "手机：15841125092 ｜ 邮箱：1482498792@qq.com ｜ 大连 ｜ 软件工程 本科", size=9.2, after=3)

    add_title(doc, "个人优势")
    for item in [
        "具备约 2 年 SAP/ERP 企业系统开发经验，熟悉需求分析、技术设计、开发测试、上线支持与线上问题排查流程。",
        "重点转向 AI 应用开发方向，熟悉 Python、FastAPI、大模型 API 调用、Prompt 设计、LangChain、RAG、向量检索等技术。",
        "能够结合企业文档、业务流程、工单记录和系统数据，开发知识库问答、文档解析、智能助手、自动化处理等 AI 应用。",
        "具备企业业务系统背景，理解数据准确性、接口稳定性、权限边界、上线变更和问题追踪等企业落地要求。",
    ]:
        add_bullet(doc, item)

    add_title(doc, "专业技能")
    skills = [
        ("AI 应用开发", "熟悉大模型 API 调用、Prompt Engineering、RAG、Embedding、文本切分、向量检索、上下文管理、流式输出等应用开发流程。"),
        ("Python 后端开发", "熟悉 Python 基础语法、面向对象、异常处理、文件处理、requests、FastAPI、RESTful API、接口封装与调试。"),
        ("知识库与检索", "熟悉 LangChain 基础用法，了解 FAISS/Chroma 等向量数据库，能够完成文档解析、切分、向量化、召回与答案生成。"),
        ("数据库与工程工具", "熟悉 SQL 基础、MySQL 基础，了解 Docker、Linux、Git、Postman、日志排查和技术文档编写。"),
        ("企业系统开发", "熟悉 SAP ABAP、Report、ALV、Smartforms、Enhancement、BAPI/RFC、IDoc、OData 等开发内容，了解 MM/SD/FI 业务流程。"),
    ]
    for title, body in skills:
        add_line(doc, f"{title}：{body}", size=9.2, after=1.3)

    add_title(doc, "项目经历")
    add_project(
        doc,
        "企业知识库智能问答系统",
        "个人项目 / 企业场景实践",
        "基于 Python + FastAPI + LangChain + 向量数据库开发企业文档问答系统，支持 PDF、Markdown、TXT 等资料解析、切分、向量化存储，并通过 RAG 实现基于知识库的问答。",
        [
            "使用 FastAPI 搭建后端接口，完成文档上传、文本解析、知识库构建、问答请求和会话记录管理。",
            "使用 LangChain 实现文档切分、Embedding 生成、向量检索和 Prompt 拼接，使回答优先基于检索资料生成。",
            "接入大模型 API 实现问答生成，支持连续问答、上下文控制和基础流式输出。",
            "设计答案来源返回机制，支持展示参考片段，提升企业知识库问答结果的可信度。",
        ],
        "Python、FastAPI、LangChain、OpenAI API/兼容大模型 API、FAISS/Chroma、HTML/CSS/JavaScript",
        "完成从文档导入、向量检索到智能问答的完整闭环，可用于企业制度查询、项目资料查询和岗位知识沉淀等场景。",
    )

    add_project(
        doc,
        "ERP/SAP 业务知识智能助手",
        "个人项目 / SAP 业务场景结合",
        "面向 SAP/ERP 开发与运维场景设计 AI 助手，将开发规范、常见问题、业务流程说明等资料构建为知识库，用于快速查询业务规则、接口说明和异常处理建议。",
        [
            "整理 SAP/ERP 相关业务文档、开发规范、问题处理记录，形成适合检索的知识库结构。",
            "基于 RAG 流程实现知识问答，使回答优先依据内部资料，减少大模型自由生成带来的不确定性。",
            "设计面向问题排查的 Prompt 模板，使回答包含可能原因、排查步骤、涉及接口和处理建议。",
            "封装问答 API，预留与网页端、企业微信或内部系统集成的接口能力。",
        ],
        "Python、FastAPI、LangChain、Embedding、向量数据库、大模型 API、SAP/ERP 业务知识",
        "将 SAP/ERP 企业系统经验转化为 AI 应用落地场景，体现 AI 在企业知识管理和运维辅助中的应用价值。",
    )

    add_project(
        doc,
        "AI 简历与文档解析工具",
        "个人项目 / 文档智能处理",
        "基于 Python 开发文档解析与信息抽取工具，支持对简历、PDF、Markdown 等文本资料进行解析、字段提取、摘要生成和改写建议输出。",
        [
            "实现文档读取、文本清洗、字段识别和结构化输出，便于后续检索或入库。",
            "通过 Prompt 模板让大模型输出岗位匹配分析、项目经历优化建议和关键词建议。",
            "封装基础命令行/接口调用方式，支持后续扩展为网页端工具。",
        ],
        "Python、文档解析、Prompt 设计、大模型 API、JSON 结构化输出",
        "完成从文档输入到结构化分析结果输出的流程，可扩展到招聘、人事、知识管理等场景。",
    )

    add_title(doc, "工作经历")
    add_role_line(doc, "松下信息系统（上海）有限公司 大连分公司", "ABAP 开发工程师 ｜ 2024.07 - 至今")
    for item in [
        "参与 SAP/ERP 系统开发与运维工作，对接业务顾问和用户需求，完成报表、表单、接口、增强类开发与问题修复。",
        "根据功能设计文档编写技术设计并完成开发、单元测试、联调测试、上线支持和运维 Ticket 处理。",
        "使用 ABAP、Open SQL、ALV、Smartforms、Enhancement、BAPI/RFC 等技术处理企业业务数据和系统集成需求。",
        "熟悉企业系统中的数据准确性、接口稳定性、上线变更和问题追踪等工程要求，为 AI 应用落地提供业务系统经验。",
    ]:
        add_bullet(doc, item)

    add_title(doc, "教育背景")
    add_role_line(doc, "大连民族大学", "软件工程 本科 ｜ 2020.09 - 2024.06")
    add_bullet(doc, "主修课程：数据结构、数据库原理、操作系统、面向对象程序设计、软件工程、计算机网络等。")
    add_bullet(doc, "英语水平：CET-4。")

    add_title(doc, "求职意向")
    add_bullet(doc, "目标岗位：AI 应用开发工程师、Python 后端开发工程师、企业智能助手/RAG 应用开发工程师。")
    add_bullet(doc, "期望方向：企业知识库问答、文档智能处理、业务流程助手、内部系统 AI 化改造、ERP/SAP 相关 AI 应用。")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
