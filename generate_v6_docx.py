# -*- coding: utf-8 -*-
"""重建 肖家伟_AI应用开发简历_v6 为结构清晰、便于编辑的 docx。
内容来源：肖家伟_AI应用开发简历_v6.pdf 的完整文本。"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = "肖家伟_AI应用开发简历_v6_可编辑.docx"
FONT = "Microsoft YaHei"
BLUE = RGBColor(44, 90, 160)
TEXT = RGBColor(34, 34, 34)


def set_run(run, size=9.2, bold=False, color=TEXT):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_line(doc, text="", size=9.2, bold=False, color=TEXT,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=1.2):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    return p


def add_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=12, bold=True, color=BLUE)
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "2C5AA0")
    border.append(bottom)
    p_pr.append(border)


def add_bullet(doc, text, size=9.2):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.first_line_indent = Cm(-0.18)
    p.paragraph_format.space_after = Pt(1.4)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run("· " + text)
    set_run(r, size=size)


def add_role_line(doc, left, right):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(left)
    set_run(r1, size=10, bold=True)
    if right:
        r2 = p.add_run("    " + right)
        set_run(r2, size=9.2, bold=True, color=RGBColor(90, 90, 90))


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9.2)
    return doc


def build():
    doc = setup_doc()

    # 头部
    add_line(doc, "肖家伟", size=22, bold=True, color=RGBColor(26, 26, 26),
             align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_line(doc, "求职方向：AI 应用开发工程师 / Python 应用开发（LLM 应用方向）",
             size=10.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_line(doc, "电话：15841125092 ｜ 邮箱：1482498792@qq.com ｜ 现居：大连 ｜ 期望工作地：不限",
             size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)

    # 核心优势
    add_title(doc, "核心优势")
    for item in [
        "AI / Python 工程实践：独立开发 SAP Smart Query Assistant，覆盖 RAG 检索、Function Calling、LangGraph Agent、自修复、Eval 评测、Web 演示与 Docker 部署。",
        "企业系统场景理解：2 年 SAP / ERP 系统一线开发经验，熟悉 SD/MM/FI 业务数据、多表关联、跨系统接口与上线验证，能把 AI 能力落到企业数据场景。",
        "求职定位清晰：目标聚焦 AI 应用开发 / Python 应用开发，优先企业知识库、智能查询、数据应用、ERP 智能化等业务方向。",
    ]:
        add_bullet(doc, item)

    # 技能清单
    add_title(doc, "技能清单")
    for item in [
        "熟悉 Python 应用开发：能够使用 Python 进行模块化工程组织、配置管理、日志与异常处理、SQL 查询、单元测试和脚本自动化；熟悉 Streamlit，了解 FastAPI / RESTful API 的接口设计方式。",
        "掌握 LLM 应用开发基础：熟悉 DeepSeek / OpenAI 兼容 API 调用、Prompt 工程、Function Calling / Tool Use、结构化输出和调用链日志分析，能够围绕业务流程设计 AI 应用链路。",
        "熟悉 RAG / Agent 工程实践：熟悉 Embedding、Chroma、Schema RAG、LangGraph 状态机、SQL 自修复和 Eval 评测；了解召回率评估、token 成本控制、工具兜底与多轮反思机制。",
        "熟悉数据库与企业数据场景：熟悉 SQL、SQLite / MySQL、多表关联、字段映射和数据清洗；理解 SAP SD/MM/FI 标准表与跨系统接口，能够把企业数据结构转化为 AI 可用的 schema 上下文。",
        "熟悉后端生态与部署：了解 Redis 基础缓存场景、Docker、Git、环境变量、容器健康检查和服务部署流程，具备从本地开发到演示部署的工程意识。",
    ]:
        add_bullet(doc, item)

    # 项目经验
    add_title(doc, "AI / Python 项目经验")
    add_role_line(doc, "SAP Smart Query Assistant · 面向 SAP 业务数据的 Text-to-SQL Agent",
                  "2026.05 - 至今 ｜ 个人项目")
    add_bullet(doc, "技术栈：Python · LangGraph · RAG · Chroma · BGE-small-zh · DeepSeek-V3 · Function Calling · SQLite · Streamlit · Docker · Eval")
    add_bullet(doc, "项目描述：面向 SAP 表名晦涩、字段语义复杂、多表关联链路长导致通用 Text-to-SQL 准确率低的问题，构建自然语言查询系统；项目可运行、可演示、可评测、可部署（可访问 http://49.232.72.207:8501）。")
    for item in [
        "工程实现：按 schema 检索、Prompt 构建、SQL 生成、schema_lookup 工具调用、SQL 执行、错误反思、结果解释等模块拆分；提供 CLI 与 Streamlit Web UI，展示 RAG 命中、tool calls、Agent attempts、最终 SQL 与查询结果。",
        "RAG 与数据：基于 SAP SD/MM/FI 场景设计 10 张核心表与约 3 万行 Faker mock 数据，使用 BGE-small-zh + Chroma 构建 Schema RAG；top-5 命中率 87.5%，平均 input token 从 1819 压缩到 1117。",
        "Agent 与工具：接入 Function Calling 设计 schema_lookup 工具，补召 MAKT/EKPO 等关键表字段；基于 LangGraph 设计 generate -> execute -> reflect 状态机，SQL 报错后携带 SQLite 错误信息重新生成，重复错误归一化与提示机制。",
        "评测与质量：自建 40 题 L1-L5 分级评测集，对比 RAG-only 与 Agent+Tools 链路，端到端通过率由 87.5% 提升至 97.5%。",
        "部署与展示：使用 Docker 构建单镜像部署，内置 SQLite 数据、Chroma 索引、Streamlit UI、密码入口与 healthcheck，便于本地复现和面试演示。",
    ]:
        add_bullet(doc, item)

    # 工作经历
    add_title(doc, "工作经历")
    add_role_line(doc, "松下信息系统（上海）有限公司 大连分公司 · ABAP 开发工程师 · 两年",
                  "2024.07 - 至今")
    for item in [
        "负责 SD/MM/FICO/PP 模块报表、接口与数据处理程序开发，参与需求澄清、方案设计、编码、测试与上线，熟悉企业系统从需求到交付的完整流程。",
        "长期处理 SAP 标准表与跨模块关联数据（如 VBAK/VBAP/KNA1/MARA/BKPF 等），具备业务数据建模、SQL / 多表关联分析和数据清洗经验。",
        "通过 RFC、服务器文件等方式对接 MES/WMS 等外部系统，理解企业系统间数据流、字段映射、异常处理与上线验证。",
    ]:
        add_bullet(doc, item)

    # 教育背景
    add_title(doc, "教育背景")
    add_role_line(doc, "大连民族大学 · 软件工程（本科）", "2020.09 - 2024.06")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
